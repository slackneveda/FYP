#!/usr/bin/env python3
"""
Automated Markdown to Thesis DOCX Converter
Converts FINAL_YEAR_PROJECT_REPORT.md to professionally formatted thesis
"""

import re
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

def setup_thesis_styles(doc):
    """Configure thesis-specific styles"""
    
    # Normal paragraph style
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    
    # Paragraph formatting
    paragraph_format = normal_style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(6)
    paragraph_format.space_before = Pt(6)
    
    # Heading 1 - Chapter titles
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_style.paragraph_format.space_before = Pt(24)
    h1_style.paragraph_format.space_after = Pt(24)
    
    # Heading 2 - Section headings
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(6)
    
    # Heading 3 - Subsection headings
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_style.paragraph_format.space_before = Pt(6)
    h3_style.paragraph_format.space_after = Pt(3)

def setup_document_margins(doc):
    """Set thesis document margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)  # Extra space for binding
        section.right_margin = Cm(2.0)

def process_title_page(doc, content_lines):
    """Process the title page section"""
    i = 0
    while i < len(content_lines):
        line = content_lines[i].strip()
        
        if line.startswith('# Sweet Dessert'):
            # Main title
            p = doc.add_paragraph(line.replace('# ', ''))
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('**THE UNIVERSITY OF LAHORE**'):
            # University name
            p = doc.add_paragraph('THE UNIVERSITY OF LAHORE')
            p.style = 'Heading 2'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('**DEPARTMENT OF TECHNOLOGY**'):
            # Department name
            p = doc.add_paragraph('DEPARTMENT OF TECHNOLOGY')
            p.style = 'Heading 2'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('**A dissertation submitted**'):
            # Dissertation statement
            p = doc.add_paragraph(line.replace('**', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'Normal'
        elif line.startswith('**Submitted by:**'):
            # Students section
            p = doc.add_paragraph('Submitted by:')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'Normal'
            # Add student names
            i += 1
            while i < len(content_lines) and not content_lines[i].startswith('**'):
                student_line = content_lines[i].strip()
                if student_line:
                    p = doc.add_paragraph(student_line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = 'Normal'
                i += 1
            continue
        elif line.startswith('**Supervised by:**'):
            # Supervisor section
            p = doc.add_paragraph('Supervised by:')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'Normal'
            i += 1
            while i < len(content_lines) and not content_lines[i].startswith('**'):
                supervisor_line = content_lines[i].strip()
                if supervisor_line:
                    p = doc.add_paragraph(supervisor_line)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.style = 'Normal'
                i += 1
            continue
        elif line.startswith('**Fall/Winter, 2025**'):
            # Date
            p = doc.add_paragraph('Fall/Winter, 2025')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'Normal'
        elif line == '---':
            # Page break
            doc.add_page_break()
        else:
            # Skip other title page content
            pass
        
        i += 1

def process_main_content(doc, content_lines):
    """Process the main thesis content"""
    i = 0
    in_declaration = False
    in_abstract = False
    
    while i < len(content_lines):
        line = content_lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle different heading levels
        if line.startswith('## Declaration'):
            p = doc.add_paragraph('DECLARATION')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_declaration = True
        elif line.startswith('## Preface'):
            p = doc.add_paragraph('PREFACE')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_declaration = False
        elif line.startswith('## Acknowledgement'):
            p = doc.add_paragraph('ACKNOWLEDGEMENTS')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_declaration = False
        elif line.startswith('## Abstract'):
            p = doc.add_paragraph('ABSTRACT')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_abstract = True
            in_declaration = False
        elif line.startswith('## LIST OF TABLES'):
            p = doc.add_paragraph('LIST OF TABLES')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            in_abstract = False
        elif line.startswith('## LIST OF FIGURES'):
            p = doc.add_paragraph('LIST OF FIGURES')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## TABLE OF CONTENTS'):
            p = doc.add_paragraph('TABLE OF CONTENTS')
            p.style = 'Heading 1'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## 1. Introduction'):
            p = doc.add_paragraph('CHAPTER 1: INTRODUCTION')
            p.style = 'Heading 1'
        elif line.startswith('## 2. Literature Review'):
            p = doc.add_paragraph('CHAPTER 2: LITERATURE REVIEW')
            p.style = 'Heading 1'
        elif line.startswith('## 3. Materials and Methods'):
            p = doc.add_paragraph('CHAPTER 3: MATERIALS AND METHODS')
            p.style = 'Heading 1'
        elif line.startswith('## 4. Results and Discussions'):
            p = doc.add_paragraph('CHAPTER 4: RESULTS AND DISCUSSIONS')
            p.style = 'Heading 1'
        elif line.startswith('## 5. Conclusions and Future Work'):
            p = doc.add_paragraph('CHAPTER 5: CONCLUSIONS AND FUTURE WORK')
            p.style = 'Heading 1'
        elif line.startswith('## 6. References'):
            p = doc.add_paragraph('REFERENCES')
            p.style = 'Heading 1'
        elif line.startswith('## Appendices'):
            p = doc.add_paragraph('APPENDICES')
            p.style = 'Heading 1'
        elif line.startswith('### '):
            # Subsection headings
            heading_text = line.replace('### ', '')
            p = doc.add_paragraph(heading_text)
            p.style = 'Heading 3'
        elif line.startswith('#### '):
            # Sub-subsection headings
            heading_text = line.replace('#### ', '')
            p = doc.add_paragraph(heading_text)
            p.style = 'Heading 3'
        elif line.startswith('**') and line.endswith('**'):
            # Bold text - convert to normal paragraph
            text = line.replace('**', '')
            p = doc.add_paragraph(text)
            p.style = 'Normal'
            if in_declaration:
                # No indentation for declaration
                p.paragraph_format.first_line_indent = 0
            elif in_abstract:
                # No indentation for abstract
                p.paragraph_format.first_line_indent = 0
            else:
                p.paragraph_format.first_line_indent = Cm(1.27)
        elif line.startswith('- **'):
            # List items in TOC
            text = line.replace('- **', '').replace('**', '')
            p = doc.add_paragraph(text)
            p.style = 'Normal'
            p.paragraph_format.first_line_indent = 0
        elif line.startswith('**Table') or line.startswith('**Figure'):
            # Table/Figure captions
            p = doc.add_paragraph(line.replace('**', ''))
            p.style = 'Normal'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = 0
        elif line == '---':
            # Page break
            doc.add_page_break()
        else:
            # Regular paragraph
            p = doc.add_paragraph(line)
            p.style = 'Normal'
            if in_declaration or in_abstract:
                # No indentation for declaration and abstract
                p.paragraph_format.first_line_indent = 0
            else:
                p.paragraph_format.first_line_indent = Cm(1.27)
        
        i += 1

def convert_markdown_to_thesis():
    """Main conversion function"""
    
    # Input and output files
    input_file = 'FINAL_YEAR_PROJECT_REPORT.md'
    output_file = 'Sweet_Dessert_Thesis_Automated.docx'
    
    # Check if input file exists
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        return
    
    # Read markdown file
    print(f"Reading {input_file}...")
    with open(input_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split into lines
    content_lines = content.split('\n')
    
    # Create Word document
    print("Creating Word document...")
    doc = Document()
    
    # Setup document formatting
    setup_document_margins(doc)
    setup_thesis_styles(doc)
    
    # Process content
    print("Processing title page...")
    process_title_page(doc, content_lines)
    
    print("Processing main content...")
    process_main_content(doc, content_lines)
    
    # Save document
    print(f"Saving thesis as {output_file}...")
    doc.save(output_file)
    
    print(f"✅ Conversion complete!")
    print(f"📄 Thesis saved as: {output_file}")
    print(f"📊 File size: {os.path.getsize(output_file) / 1024:.1f} KB")
    
    return output_file

if __name__ == "__main__":
    try:
        output_file = convert_markdown_to_thesis()
        print(f"\n🎉 Your thesis is ready!")
        print(f"📁 You can find it here: {os.path.abspath(output_file)}")
        print(f"\n📋 Next steps:")
        print(f"1. Open the DOCX file in Microsoft Word")
        print(f"2. Add page numbers (Roman for preliminaries, Arabic for main content)")
        print(f"3. Generate automatic Table of Contents if needed")
        print(f"4. Review and adjust any formatting as needed")
        
    except Exception as e:
        print(f"❌ Error during conversion: {str(e)}")
        print("Please check your input file and try again.")
